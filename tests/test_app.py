import io
from datetime import datetime, timezone
from pathlib import Path

import numpy as np
import pandas as pd
import pytest
from PIL import Image
from docx import Document

import app


@pytest.fixture(autouse=True)
def isolated_storage(tmp_path, monkeypatch):
    base_dir = tmp_path / "app_root"
    data_dir = base_dir / "data"
    events_dir = data_dir / "events"
    base_dir.mkdir()

    monkeypatch.setattr(app, "BASE_DIR", base_dir)
    monkeypatch.setattr(app, "DATA_DIR", data_dir)
    monkeypatch.setattr(app, "EVENTS_DIR", events_dir)
    assets_dir = base_dir / "assets"
    assets_dir.mkdir()
    monkeypatch.setattr(app, "ASSETS_DIR", assets_dir)
    monkeypatch.setattr(app, "EVENTS_FILE", data_dir / "events.csv")
    monkeypatch.setattr(app, "LEGACY_ATTENDEE_FILE", data_dir / "attendees.csv")
    monkeypatch.setattr(app, "LEGACY_SIGNIN_FILE", data_dir / "signins.csv")
    monkeypatch.setattr(app, "LEGACY_SIGNATURE_DIR", data_dir / "signatures")

    for cached in (app.load_events, app.load_attendees, app.load_signins):
        try:
            cached.clear()
        except Exception:  # pragma: no cover - streamlit runtime mismatch
            pass

    app.ensure_storage()

    template_doc = Document()
    table = template_doc.add_table(rows=14, cols=4)
    header = table.rows[0].cells
    header[0].text = "#"
    header[1].text = "Name and Surname"
    header[2].text = "Organization"
    header[3].text = "Signature"
    for idx in range(1, 14):
        table.rows[idx].cells[0].text = str(idx)
    template_doc.save(assets_dir / "SignatureList.docx")

    front_doc = Document()
    front_doc.add_heading("INNO2MARE Cover", level=1)
    front_doc.add_paragraph("EDUKACIJA: EDUCATION")
    info_table = front_doc.add_table(rows=3, cols=2)
    info_table.autofit = True
    info_table.rows[0].cells[0].text = "Datum:"
    info_table.rows[0].cells[1].text = "Date"
    info_table.rows[1].cells[0].text = "Lokacija:"
    info_table.rows[1].cells[1].text = "Location"
    info_table.rows[2].cells[0].text = "Projektna aktivnost:"
    info_table.rows[2].cells[1].text = "Activity"
    front_doc.save(assets_dir / "INNO2MARE_FrontPage.docx")
    yield

    for cached in (app.load_events, app.load_attendees, app.load_signins):
        try:
            cached.clear()
        except Exception:  # pragma: no cover
            pass


def read_csv(path: Path) -> pd.DataFrame:
    return pd.read_csv(path, dtype=str).fillna("")


def test_ensure_storage_creates_default_event():
    assert app.DATA_DIR.exists()
    assert app.EVENTS_DIR.exists()
    assert app.EVENTS_FILE.exists()

    events = read_csv(app.EVENTS_FILE)
    assert list(events.columns) == app.EVENT_COLUMNS
    assert app.DEFAULT_EVENT_ID in events["event_id"].values
    default_row = events.loc[events["event_id"] == app.DEFAULT_EVENT_ID].iloc[0]
    assert default_row["declaration"] == app.DECLARATION_PLACEHOLDER
    assert default_row["project_type"] == app.PROJECT_TYPES[0]
    assert default_row["is_default"] == "true"

    default_attendees = read_csv(app.attendee_file(app.DEFAULT_EVENT_ID))
    default_signins = read_csv(app.signin_file(app.DEFAULT_EVENT_ID))
    assert list(default_attendees.columns) == app.ATTENDEE_COLUMNS
    assert list(default_signins.columns) == app.SIGNIN_COLUMNS


def test_replace_attendees_validates_required_columns():
    invalid = pd.DataFrame(
        [
            {
                "attendee_id": "1",
                "name": "Jane",
                "company": "ACME",
                "email": "jane@acme.com",
            },
        ]
    )
    with pytest.raises(ValueError):
        app.replace_attendees(app.DEFAULT_EVENT_ID, invalid)


def test_replace_attendees_overwrites_existing_data():
    initial = pd.DataFrame(
        [
            {
                "attendee_id": "1",
                "name": "Jane Doe",
                "company": "ACME",
                "email": "jane@acme.com",
                "phone": "+385-91-555-0000",
            }
        ]
    )
    app.replace_attendees(app.DEFAULT_EVENT_ID, initial)

    updated = pd.DataFrame(
        [
            {
                "attendee_id": "2",
                "name": "John Smith",
                "company": "Widgets",
                "email": "john@widgets.com",
                "phone": "+385-91-555-0001",
            }
        ]
    )
    app.replace_attendees(app.DEFAULT_EVENT_ID, updated)

    stored = read_csv(app.attendee_file(app.DEFAULT_EVENT_ID))
    assert len(stored) == 1
    assert stored.iloc[0]["attendee_id"] == "2"


def test_create_event_generates_unique_directory():
    record = app.create_event(
        name="Advanced Session",
        date="2024-11-01",
        location="Main HQ",
        project_activity="Advanced Module",
        project_type="GREENPACT",
        declaration="Custom declaration text",
        description="Hands-on workshop",
    )
    event_id = record["event_id"]
    assert (app.EVENTS_DIR / event_id).exists()
    events = read_csv(app.EVENTS_FILE)
    assert events.shape[0] == 2
    created_row = events.loc[events["event_id"] == event_id].iloc[0]
    assert created_row["date"] == "2024-11-01"
    assert created_row["declaration"] == "Custom declaration text"
    assert created_row["project_type"] == "GREENPACT"
    assert created_row["is_default"] == "false"


def test_update_event_details_overwrites_fields():
    record = app.create_event(
        name="Session A",
        date="2024-10-01",
        location="City A",
        project_activity="Module A",
        project_type="EDIH",
        declaration="Declaration A",
        description="Notes A",
    )
    app.update_event_details(
        record["event_id"],
        name="Session B",
        date="2024-10-02",
        location="City B",
        project_activity="Module B",
        project_type="EEN",
        declaration="Declaration B",
        description="Notes B",
    )
    events = read_csv(app.EVENTS_FILE)
    row = events.loc[events["event_id"] == record["event_id"]].iloc[0]
    assert row["name"] == "Session B"
    assert row["location"] == "City B"
    assert row["project_activity"] == "Module B"
    assert row["project_type"] == "EEN"
    assert row["is_default"] == "false"


def test_set_default_event_updates_flags():
    record = app.create_event(
        name="Session Default",
        date="2024-12-01",
        location="Main Hall",
        project_activity="Workshop",
        project_type="INNO2MARE",
        declaration="Custom declaration text",
        description="Notes",
    )
    app.set_default_event(record["event_id"])
    events = read_csv(app.EVENTS_FILE)
    defaults = events[events["is_default"] == "true"]
    assert len(defaults) == 1
    assert defaults.iloc[0]["event_id"] == record["event_id"]


def test_filter_attendees_matches_by_partial_text():
    df = pd.DataFrame(
        [
            {
                "attendee_id": "1",
                "name": "Alice Example",
                "company": "Example Co",
                "email": "alice@example.com",
                "phone": "+385-91-555-0101",
            },
            {
                "attendee_id": "2",
                "name": "Bob Builder",
                "company": "Construction Ltd",
                "email": "bob@build.com",
                "phone": "+385-91-555-0102",
            },
        ]
    )

    result_name = app.filter_attendees(df, "alice")
    assert list(result_name["attendee_id"]) == ["1"]

    result_company = app.filter_attendees(df, "construction")
    assert list(result_company["attendee_id"]) == ["2"]

    result_email = app.filter_attendees(df, "@example.com")
    assert list(result_email["attendee_id"]) == ["1"]

    result_empty_query = app.filter_attendees(df, "")
    pd.testing.assert_frame_equal(result_empty_query.reset_index(drop=True), df)


def test_is_signature_blank_detects_drawn_pixels():
    blank = np.full((5, 5, 4), 255, dtype=np.uint8)
    assert app.is_signature_blank(blank)

    drawn = blank.copy()
    drawn[2, 2, :3] = 0
    assert not app.is_signature_blank(drawn)


def test_save_signature_image_persists_file():
    event_id = app.DEFAULT_EVENT_ID
    pixels = np.full((10, 10, 4), 255, dtype=np.uint8)
    pixels[3, 4, :3] = 0  # simulate a stroke

    relative_path = app.save_signature_image(event_id, pixels)
    stored_path = app.BASE_DIR / relative_path

    assert relative_path.startswith(f"data/events/{event_id}/signatures/")
    assert stored_path.exists()

    with Image.open(stored_path) as img:
        assert img.mode == "RGB"
        assert img.size == (10, 10)


def test_append_signin_appends_to_event_csv():
    event_id = app.DEFAULT_EVENT_ID
    entry = {
        "record_id": "abc123",
        "attendee_id": "1",
        "name": "Alice Example",
        "company": "Example Co",
        "phone": "+385-91-555-0201",
        "email": "alice@example.com",
        "signed_at": "2024-01-01T12:00:00Z",
        "signature_file": f"data/events/{event_id}/signatures/test.png",
    }

    app.append_signin(event_id, entry)

    stored = read_csv(app.signin_file(event_id))
    assert len(stored) == 1
    assert stored.iloc[0]["record_id"] == "abc123"


def test_generate_signature_document_returns_bytes():
    event_id = app.DEFAULT_EVENT_ID
    entry = {
        "record_id": "xyz789",
        "attendee_id": "3",
        "name": "Bob Signed",
        "company": "SignCo",
        "phone": "+385-91-555-0301",
        "email": "bob@sign.co",
        "signed_at": "2024-05-01T09:00:00Z",
        "signature_file": "",
    }
    app.append_signin(event_id, entry)
    filename, data = app.generate_signature_document(event_id)
    assert filename.endswith(".docx")
    assert len(data) > 0


def test_generate_signature_document_inno2mare_front_page_and_signature_image():
    event = app.create_event(
        name="INNO2MARE Session",
        date="2024-07-15",
        location="Rijeka Hub",
        project_activity="Innovation Workshop",
        project_type="INNO2MARE",
        declaration="Custom declaration",
        description="",
    )
    event_id = event["event_id"]

    signature_pixels = np.full((12, 12, 4), 255, dtype=np.uint8)
    signature_pixels[3:9, 3:9, :3] = 0
    signature_path = app.save_signature_image(event_id, signature_pixels)

    record_id = "sig-001"
    app.append_signin(
        event_id,
        {
            "record_id": record_id,
            "attendee_id": "42",
            "name": "Marina Innovator",
            "company": "BlueTech",
            "phone": "+385-91-555-0401",
            "email": "marina@bluetech.hr",
            "signed_at": datetime.now(timezone.utc).isoformat(),
            "signature_file": signature_path,
        },
    )

    filename, data = app.generate_signature_document(event_id)
    assert filename == f"{event_id}-signature-list.docx"

    document = Document(io.BytesIO(data))
    combined_text_parts = []
    combined_text_parts.extend(paragraph.text for paragraph in document.paragraphs)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                combined_text_parts.extend(p.text for p in cell.paragraphs)
    combined_text = "\n".join(filter(None, combined_text_parts))
    assert "INNO2MARE Session" in combined_text
    assert "Rijeka Hub" in combined_text
    assert "Innovation Workshop" in combined_text
    assert "<w:br w:type=\"page\"" not in document._element.xml

    four_column_table = next(table for table in document.tables if len(table.rows[0].cells) == 4)

    first_data_row = four_column_table.rows[1]
    assert first_data_row.cells[1].text == "Marina Innovator"
    assert first_data_row.cells[2].text == "BlueTech"
    signature_cell_xml = first_data_row.cells[3]._tc.xml
    assert "blip" in signature_cell_xml
