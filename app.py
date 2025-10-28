from __future__ import annotations

import uuid
from datetime import datetime, timezone
import os
import re
import shutil
from pathlib import Path
import io

import numpy as np
import pandas as pd
import streamlit as st
from PIL import Image
from streamlit_drawable_canvas import st_canvas
from docx import Document
from docx.shared import Inches


BASE_DIR = Path(__file__).parent
DATA_DIR = BASE_DIR / "data"
EVENTS_DIR = DATA_DIR / "events"
EVENTS_FILE = DATA_DIR / "events.csv"
ASSETS_DIR = BASE_DIR / "assets"
LEGACY_ATTENDEE_FILE = DATA_DIR / "attendees.csv"
LEGACY_SIGNIN_FILE = DATA_DIR / "signins.csv"
LEGACY_SIGNATURE_DIR = DATA_DIR / "signatures"

EVENT_COLUMNS = [
    "event_id",
    "name",
    "date",
    "location",
    "project_activity",
    "project_type",
    "is_default",
    "declaration",
    "description",
]
ATTENDEE_COLUMNS = ["attendee_id", "name", "company", "email"]
SIGNIN_COLUMNS = [
    "record_id",
    "attendee_id",
    "name",
    "company",
    "email",
    "signed_at",
    "signature_file",
]

ADMIN_USERNAME = os.getenv("ADMIN_USERNAME", "step")
ADMIN_PASSWORD = os.getenv("ADMIN_PASSWORD", "step")

DEFAULT_EVENT_ID = "default"
DEFAULT_EVENT_NAME = "Education Session"
DECLARATION_PLACEHOLDER = (
    "NAPOMENA: Potpisom na potpisnoj listi dajem svoju izričitu suglasnost i privolu "
    "da STEP RI prikuplja, obrađuje i koristi moje osobne podatke u svrhe evidentiranja "
    "broja polaznika kao dokaz provedene aktivnosti; isti mogu biti ustupljeni trećim "
    "stranama kao kontrolnom tijelu, ukoliko je aktivnost organizirana u sklopu EU i "
    "međunarodnih projekata.\n\n"
    "NOTE: By signing on the signature list I give my consent and approval for STEP RI "
    "to collect, process and use my personal data for the purposes of documenting the "
    "number of attendants as evidence of the activity carried out; it may also be "
    "provided to third parties presenting a control body, if the activity is organized "
    "as a part of the EU and international projects."
)

PROJECT_TYPES = ["INNO2MARE", "EDIH", "EEN", "GREENPACT"]
INNO2MARE_LOGO_WIDTH = 220
PROJECT_TEMPLATES = {
    "INNO2MARE": {
        "tagline": "INNO2MARE project – 101087348 – funded by Horizon Europe",
        "description": "Empowering maritime regions through innovation and collaboration.",
        "image": "assets/eu-funded.png",
    },
    "EDIH": {
        "tagline": "European Digital Innovation Hub activities",
        "description": "Supporting SMEs on their digital transformation journey.",
        "image": "assets/image.png",
    },
    "EEN": {
        "tagline": "Enterprise Europe Network initiatives",
        "description": "Connecting businesses to grow on an international scale.",
        "image": "assets/image.png",
    },
    "GREENPACT": {
        "tagline": "GREENPACT sustainability programme",
        "description": "Advancing green and sustainable business practices.",
        "image": "assets/image.png",
    },
}


def event_dir(event_id: str) -> Path:
    return EVENTS_DIR / event_id


def attendee_file(event_id: str) -> Path:
    return event_dir(event_id) / "attendees.csv"


def signin_file(event_id: str) -> Path:
    return event_dir(event_id) / "signins.csv"


def signature_dir(event_id: str) -> Path:
    return event_dir(event_id) / "signatures"


def slugify(value: str) -> str:
    base = re.sub(r"[^a-z0-9]+", "-", (value or "").lower()).strip("-")
    return base or uuid.uuid4().hex[:8]


def get_project_template(project_type: str) -> dict[str, str]:
    return PROJECT_TEMPLATES.get(project_type, PROJECT_TEMPLATES[PROJECT_TYPES[0]])


def filter_attendees(attendees: pd.DataFrame, query: str) -> pd.DataFrame:
    """Return attendees whose name, company, or email contains the query string."""
    cleaned_query = (query or "").strip().lower()
    if not cleaned_query:
        return attendees
    searchable = (
        attendees[["name", "company", "email"]]
        .astype(str)
        .agg(" ".join, axis=1)
        .str.lower()
    )
    mask = searchable.str.contains(cleaned_query, regex=False)
    return attendees[mask]


def ensure_storage() -> None:
    """Create required data folders/files if they do not yet exist."""
    DATA_DIR.mkdir(parents=True, exist_ok=True)
    EVENTS_DIR.mkdir(parents=True, exist_ok=True)
    ASSETS_DIR.mkdir(parents=True, exist_ok=True)

    if not EVENTS_FILE.exists():
        pd.DataFrame(
            [
                {
                    "event_id": DEFAULT_EVENT_ID,
                    "name": DEFAULT_EVENT_NAME,
                    "date": "",
                    "location": "",
                    "project_activity": "",
                    "project_type": PROJECT_TYPES[0],
                    "is_default": "true",
                    "declaration": DECLARATION_PLACEHOLDER,
                    "description": "",
                }
            ],
            columns=EVENT_COLUMNS,
        ).to_csv(EVENTS_FILE, index=False)

    events_df = pd.read_csv(EVENTS_FILE, dtype=str).fillna("")
    missing = [col for col in EVENT_COLUMNS if col not in events_df.columns]
    if missing:
        defaults = {
            "declaration": DECLARATION_PLACEHOLDER,
            "project_type": PROJECT_TYPES[0],
            "is_default": "false",
        }
        for col in missing:
            events_df[col] = defaults.get(col, "")
    events_df = events_df[EVENT_COLUMNS]
    events_df.to_csv(EVENTS_FILE, index=False)

    for event_id in events_df["event_id"]:
        ensure_event_storage(event_id)

    # Legacy data migration (single-session layout -> default event)
    default_attendee_path = attendee_file(DEFAULT_EVENT_ID)
    default_signin_path = signin_file(DEFAULT_EVENT_ID)
    default_signature_path = signature_dir(DEFAULT_EVENT_ID)

    if LEGACY_ATTENDEE_FILE.exists() and not default_attendee_path.exists():
        shutil.copy(LEGACY_ATTENDEE_FILE, default_attendee_path)
    if LEGACY_SIGNIN_FILE.exists() and not default_signin_path.exists():
        shutil.copy(LEGACY_SIGNIN_FILE, default_signin_path)
    if LEGACY_SIGNATURE_DIR.exists() and not any(default_signature_path.iterdir()):
        for legacy_file in LEGACY_SIGNATURE_DIR.glob("*.png"):
            shutil.copy(legacy_file, default_signature_path / legacy_file.name)


def ensure_event_storage(event_id: str) -> None:
    path = event_dir(event_id)
    path.mkdir(parents=True, exist_ok=True)

    attendees_path = attendee_file(event_id)
    signins_path = signin_file(event_id)
    signatures_path = signature_dir(event_id)

    if not attendees_path.exists():
        pd.DataFrame(columns=ATTENDEE_COLUMNS).to_csv(attendees_path, index=False)
    if not signins_path.exists():
        pd.DataFrame(columns=SIGNIN_COLUMNS).to_csv(signins_path, index=False)
    signatures_path.mkdir(parents=True, exist_ok=True)


@st.cache_data(show_spinner=False)
def load_events() -> pd.DataFrame:
    if not EVENTS_FILE.exists():
        ensure_storage()
    df = pd.read_csv(EVENTS_FILE, dtype=str).fillna("")
    missing = [col for col in EVENT_COLUMNS if col not in df.columns]
    if missing:
        raise ValueError(f"Events file missing required columns: {', '.join(missing)}")
    return df[EVENT_COLUMNS]


@st.cache_data(show_spinner=False)
def load_attendees(event_id: str) -> pd.DataFrame:
    file_path = attendee_file(event_id)
    if not file_path.exists():
        return pd.DataFrame(columns=ATTENDEE_COLUMNS)
    df = pd.read_csv(file_path, dtype=str).fillna("")
    missing = [col for col in ATTENDEE_COLUMNS if col not in df.columns]
    if missing:
        raise ValueError(
            f"Attendee CSV missing required columns: {', '.join(missing)}"
        )
    return df[ATTENDEE_COLUMNS]


@st.cache_data(show_spinner=False)
def load_signins(event_id: str) -> pd.DataFrame:
    file_path = signin_file(event_id)
    if not file_path.exists():
        return pd.DataFrame(columns=SIGNIN_COLUMNS)
    df = pd.read_csv(file_path, dtype=str).fillna("")
    missing = [col for col in SIGNIN_COLUMNS if col not in df.columns]
    if missing:
        raise ValueError(f"Sign-in CSV missing required columns: {', '.join(missing)}")
    return df[SIGNIN_COLUMNS]


def write_events(df: pd.DataFrame) -> None:
    df = df.astype(str).fillna("")
    df = df[EVENT_COLUMNS]
    df.to_csv(EVENTS_FILE, index=False)
    load_events.clear()


def create_event(
    name: str,
    date: str = "",
    location: str = "",
    project_activity: str = "",
    project_type: str = PROJECT_TYPES[0],
    declaration: str = "",
    description: str = "",
) -> dict[str, str]:
    events = load_events()
    existing_ids = set(events["event_id"])

    base_slug = slugify(name)
    candidate = base_slug or DEFAULT_EVENT_ID
    suffix = 1
    while candidate in existing_ids:
        candidate = f"{base_slug}-{suffix}"
        suffix += 1

    event_id = candidate
    record = {
        "event_id": event_id,
        "name": name or f"Session {len(events) + 1}",
        "date": date,
        "location": location,
        "project_activity": project_activity,
        "project_type": project_type if project_type in PROJECT_TYPES else PROJECT_TYPES[0],
        "is_default": "false",
        "declaration": declaration or DECLARATION_PLACEHOLDER,
        "description": description,
    }
    updated = pd.concat([events, pd.DataFrame([record])], ignore_index=True)
    write_events(updated)
    ensure_event_storage(event_id)
    return record


def update_event_details(
    event_id: str,
    name: str,
    date: str,
    location: str,
    project_activity: str,
    project_type: str,
    declaration: str,
    description: str,
) -> None:
    events = load_events()
    if event_id not in events["event_id"].values:
        raise ValueError(f"Event '{event_id}' does not exist.")
    updated_declaration = declaration or DECLARATION_PLACEHOLDER
    events.loc[
        events["event_id"] == event_id,
        [
            "name",
            "date",
            "location",
            "project_activity",
            "project_type",
            "declaration",
            "description",
        ],
    ] = [
        name,
        date,
        location,
        project_activity,
        project_type if project_type in PROJECT_TYPES else PROJECT_TYPES[0],
        updated_declaration,
        description,
    ]
    write_events(events)


def set_default_event(event_id: str) -> None:
    events = load_events()
    if event_id not in events["event_id"].values:
        raise ValueError(f"Event '{event_id}' does not exist.")
    events["is_default"] = events["event_id"].apply(
        lambda eid: "true" if eid == event_id else "false"
    )
    write_events(events)


def append_signin(event_id: str, entry: dict[str, str]) -> None:
    df = load_signins(event_id)
    df = pd.concat([df, pd.DataFrame([entry])], ignore_index=True)
    df.to_csv(signin_file(event_id), index=False)
    load_signins.clear()


def generate_signature_document(event_id: str) -> tuple[str, bytes]:
    template_path = ASSETS_DIR / "SignatureList.docx"
    if not template_path.exists():
        raise FileNotFoundError("Signature template not found in assets directory.")

    signins = load_signins(event_id).copy()
    if signins.empty:
        raise ValueError("No sign-in records available for this session.")

    if "signed_at" in signins.columns:
        signins = signins.sort_values(
            by="signed_at", ascending=True, na_position="last"
        )

    document = Document(template_path)
    table = document.tables[0]

    header_offset = 1
    required_rows = header_offset + len(signins)
    while len(table.rows) < required_rows:
        table.add_row()

    for idx, (_, row) in enumerate(signins.iterrows(), start=1):
        table_row = table.rows[idx]
        table_row.cells[0].text = str(idx)
        table_row.cells[1].text = row.get("name", "")
        table_row.cells[2].text = row.get("company", "")

        signature_cell = table_row.cells[3]
        signature_cell.text = ""
        signature_path_str = row.get("signature_file", "")
        signature_path = BASE_DIR / signature_path_str if signature_path_str else None
        if signature_path and signature_path.exists():
            run = signature_cell.paragraphs[0].add_run()
            run.add_picture(str(signature_path), width=Inches(1.5))

    for idx in range(required_rows, len(table.rows)):
        for cell in table.rows[idx].cells:
            cell.text = ""

    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    filename = f"{event_id}-signature-list.docx"
    return filename, buffer.getvalue()


def replace_attendees(event_id: str, new_df: pd.DataFrame) -> None:
    new_df = new_df.astype(str).fillna("")
    missing = [col for col in ATTENDEE_COLUMNS if col not in new_df.columns]
    if missing:
        raise ValueError(
            f"Uploaded attendee list missing required columns: {', '.join(missing)}"
        )
    new_df.to_csv(attendee_file(event_id), index=False)
    load_attendees.clear()


def save_signature_image(event_id: str, image_data: np.ndarray) -> str:
    """Persist the signature image and return its relative path."""
    if image_data is None:
        raise ValueError("No signature data to save.")
    image = Image.fromarray((image_data).astype("uint8"))
    # Convert transparent pixels to white background for readability.
    background = Image.new("RGBA", image.size, "WHITE")
    composed = Image.alpha_composite(background, image).convert("RGB")

    file_id = uuid.uuid4().hex
    signature_path = signature_dir(event_id)
    signature_path.mkdir(parents=True, exist_ok=True)
    file_path = signature_path / f"{file_id}.png"
    composed.save(file_path, format="PNG")
    relative_path = file_path.relative_to(BASE_DIR)
    return relative_path.as_posix()


def is_signature_blank(image_data: np.ndarray | None) -> bool:
    if image_data is None:
        return True
    # Check if all RGB pixels are white (255). If so, no stroke was drawn.
    rgb_pixels = image_data[:, :, :3]
    return np.all(rgb_pixels == 255)


def sign_in_page(event_id: str, event: dict[str, str]) -> None:
    event_name = event.get("name") or event_id
    st.header(f"Education Sign-In - {event_name}")
    attendees = load_attendees(event_id)

    recent_success = st.session_state.pop("sign_in_success", None)
    if isinstance(recent_success, dict):
        st.success(
            f"Thank you, {recent_success.get('name')}! "
            f"You are signed in for {recent_success.get('event', event_name)}."
        )
    elif recent_success:
        st.success(f"Thank you, {recent_success}! You are signed in for {event_name}.")

    detail_col1, detail_col2 = st.columns(2)
    detail_col1.markdown(
        f"**Education date:** {event.get('date') or 'To be confirmed'}"
    )
    detail_col1.markdown(
        f"**Location:** {event.get('location') or 'To be confirmed'}"
    )
    detail_col2.markdown(
        f"**Project activity:** {event.get('project_activity') or 'Not specified'}"
    )

    declaration_text = event.get("declaration") or DECLARATION_PLACEHOLDER
    project_type = event.get("project_type", "")
    template = get_project_template(project_type)

    def render_footer() -> None:
        st.divider()
        st.markdown("**Declaration**")
        st.info(declaration_text)
        if project_type == "INNO2MARE":
            if template.get("tagline"):
                st.markdown(f"**{template['tagline']}**")
            if template.get("description"):
                st.caption(template["description"])
            image_path = BASE_DIR / template.get("image", "")
            if image_path.exists():
                st.image(image_path, width=INNO2MARE_LOGO_WIDTH)
            else:
                st.warning(
                    f"Project image '{template.get('image')}' is missing. "
                    "Upload the asset under the repository's assets folder."
                )
            st.caption("Funded by the European Union")

    if attendees.empty:
        st.warning(
            "No attendee list found for this session. Please contact the admin to upload one."
        )
        render_footer()
        return

    use_existing = st.toggle(
        "I'm on the attendee list", value=True, help="Disable if you are a walk-in."
    )

    selected_attendee = None
    attendee_id = ""
    name = ""
    company = ""
    email = ""

    if use_existing:
        if st.session_state.pop("reset_attendee_search", False):
            st.session_state.pop("attendee_search", None)

        selected_attendee = st.session_state.get("selected_attendee")
        search_query = st.text_input(
            "Search for your name, company, or email",
            key="attendee_search",
            placeholder="Start typing to filter the list...",
        )

        filtered_attendees = filter_attendees(attendees, search_query)
        filtered_records = filtered_attendees.to_dict("records")

        if selected_attendee:

            def matches_selected(record: dict[str, str]) -> bool:
                record_id = (record.get("attendee_id") or "").strip()
                selected_id = (selected_attendee.get("attendee_id") or "").strip()
                if record_id and selected_id:
                    return record_id == selected_id
                return (
                    record.get("name") == selected_attendee.get("name")
                    and record.get("email") == selected_attendee.get("email")
                )

            if not any(matches_selected(record) for record in filtered_records):
                selected_attendee = None
                st.session_state.pop("selected_attendee", None)

        max_results = 12
        if len(filtered_records) > max_results:
            st.info(
                f"Showing first {max_results} matches. Refine your search for more precise results."
            )
            filtered_records = filtered_records[:max_results]

        if not filtered_records:
            st.warning(
                "No attendees match that search. Try a different name or switch off the toggle if you are a walk-in."
            )
        else:
            st.caption("Tap your name to fill in the form automatically.")
            columns_per_row = 2
            for start in range(0, len(filtered_records), columns_per_row):
                row_records = filtered_records[start : start + columns_per_row]
                row_cols = st.columns(len(row_records))
                for idx, (col, record) in enumerate(zip(row_cols, row_records)):
                    name_line = record.get("name") or "Unnamed attendee"
                    button_label = name_line
                    unique_fragment = (
                        record.get("attendee_id")
                        or record.get("email")
                        or record.get("name")
                        or f"{start}_{idx}"
                    )
                    button_key = f"attendee_btn_{unique_fragment}_{start}_{idx}"
                    if col.button(
                        button_label, key=button_key, use_container_width=True
                    ):
                        st.session_state["selected_attendee"] = record
                        selected_attendee = record

        selected_attendee = st.session_state.get("selected_attendee")
        if selected_attendee:
            st.info(f"Selected: {selected_attendee.get('name', 'Unknown attendee')}")
            attendee_id = selected_attendee.get("attendee_id", "")
            name = selected_attendee.get("name", "")
            company = selected_attendee.get("company", "")
            email = selected_attendee.get("email", "")
    else:
        st.session_state.pop("selected_attendee", None)
        st.session_state.pop("attendee_search", None)

    name = st.text_input("Full name*", value=name)
    company = st.text_input("Company", value=company)
    email = st.text_input("Email", value=email)

    st.markdown("#### Signature")
    st.caption("Please sign inside the box below.")

    if "signature_canvas_key" not in st.session_state:
        st.session_state["signature_canvas_key"] = f"signature_canvas_{uuid.uuid4().hex}"
    canvas_key = st.session_state["signature_canvas_key"]

    canvas_result = st_canvas(
        fill_color="#FFFFFF",
        stroke_width=2,
        stroke_color="#000000",
        background_color="#FFFFFF",
        height=200,
        width=600,
        drawing_mode="freedraw",
        key=canvas_key,
    )

    submitted = st.button("Submit sign-in", type="primary")

    if submitted:
        errors = []
        if not name.strip():
            errors.append("Name is required.")
        if is_signature_blank(getattr(canvas_result, "image_data", None)):
            errors.append("Signature is required.")

        if errors:
            for err in errors:
                st.error(err)
            return

        try:
            signature_path = save_signature_image(event_id, canvas_result.image_data)
        except ValueError as exc:
            st.error(str(exc))
            return

        entry = {
            "record_id": uuid.uuid4().hex,
            "attendee_id": attendee_id,
            "name": name.strip(),
            "company": company.strip(),
            "email": email.strip(),
            "signed_at": datetime.now(timezone.utc).isoformat(),
            "signature_file": signature_path,
        }
        append_signin(event_id, entry)
        st.session_state.pop("selected_attendee", None)
        st.session_state["sign_in_success"] = {
            "name": entry["name"],
            "event": event_name,
        }
        st.session_state["reset_attendee_search"] = True
        st.session_state["signature_canvas_key"] = (
            f"signature_canvas_{uuid.uuid4().hex}"
        )
        st.rerun()

    render_footer()


def admin_login_page() -> None:
    st.header("Admin Login")
    st.caption(
        "Enter the admin credentials. You can override them with the "
        "`ADMIN_USERNAME` and `ADMIN_PASSWORD` environment variables."
    )

    with st.form("admin_login"):
        username = st.text_input("Username")
        password = st.text_input("Password", type="password")
        submitted = st.form_submit_button("Log in")

    if submitted:
        if username == ADMIN_USERNAME and password == ADMIN_PASSWORD:
            st.session_state["admin_authenticated"] = True
            st.success("Logged in successfully.")
            st.rerun()
        else:
            st.error("Invalid credentials. Please try again.")


def admin_page(
    event_id: str, events_df: pd.DataFrame, active_event: dict[str, str]
) -> None:
    st.header(f"Admin Panel - {active_event.get('name') or event_id}")
    st.caption(
        "Manage attendee lists, monitor sign-ins, and download data for certificates."
    )

    events_df = load_events()
    event_ids = list(events_df["event_id"])
    if event_id not in event_ids:
        event_id = event_ids[0]
        st.session_state["active_event_id"] = event_id

    labels_by_id = {}
    for _, row in events_df.iterrows():
        label = f"{row['name'] or row['event_id']} - {row['date'] or 'Date TBD'}"
        if str(row.get("is_default", "")).lower() == "true":
            label += " (default)"
        labels_by_id[row["event_id"]] = label
    default_index = event_ids.index(event_id)

    st.subheader("Education sessions")
    selected_for_edit = st.radio(
        "Choose a session to edit",
        event_ids,
        index=default_index,
        format_func=lambda eid: labels_by_id[eid],
        key="admin_event_radio",
    )
    if selected_for_edit != event_id:
        st.session_state["active_event_id"] = selected_for_edit
        st.session_state["reset_attendee_search"] = True
        st.session_state.pop("selected_attendee", None)
        st.session_state.pop("signature_canvas_key", None)
        st.rerun()

    event_id = selected_for_edit
    active_event = events_df.set_index("event_id").loc[event_id].to_dict()
    is_default = str(active_event.get("is_default", "")).lower() == "true"
    if is_default:
        st.caption("This session is currently the default selection on the sign-in page.")
    elif st.button("Make this the default session", key=f"make_default_{event_id}"):
        try:
            set_default_event(event_id)
            st.success("Default session updated.")
            st.session_state["active_event_id"] = event_id
            st.session_state.pop("editing_event_id", None)
            st.rerun()
        except Exception as exc:  # noqa: BLE001
            st.error(f"Unable to update default session: {exc}")

    attendees = load_attendees(event_id)
    signins = load_signins(event_id)

    col1, col2 = st.columns(2)
    col1.metric("Registered attendees", len(attendees))
    col2.metric("Sign-ins collected", len(signins))

    with st.expander("Current attendee list", expanded=False):
        st.dataframe(attendees)

    with st.expander("Sign-in records", expanded=False):
        st.dataframe(signins)

    editing_event = st.session_state.get("editing_event_id")
    if editing_event not in event_ids:
        st.session_state.pop("editing_event_id", None)
        editing_event = None

    if editing_event == event_id:
        st.subheader("Session details")
        with st.form("update_event_details"):
            current_project_type = active_event.get("project_type", PROJECT_TYPES[0])
            try:
                default_project_index = PROJECT_TYPES.index(current_project_type)
            except ValueError:
                default_project_index = 0
            updated_project_type = st.selectbox(
                "Project template*",
                PROJECT_TYPES,
                index=default_project_index,
                help="Select the project this education belongs to. This controls the banner and template.",
            )
            template_preview = get_project_template(updated_project_type)
            st.caption(
                f"Preview image: {template_preview.get('image', 'N/A')} — "
                f"{template_preview.get('tagline', '') or 'No tagline'}"
            )
            updated_name = st.text_input(
                "Education name*",
                value=active_event.get("name", ""),
                max_chars=200,
            )
            col_date, col_location = st.columns(2)
            updated_date = col_date.text_input(
                "Education date*",
                value=active_event.get("date", ""),
                placeholder="YYYY-MM-DD",
            )
            updated_location = col_location.text_input(
                "Location*",
                value=active_event.get("location", ""),
                placeholder="City / venue",
                max_chars=200,
            )
            updated_project_activity = st.text_input(
                "Project activity*",
                value=active_event.get("project_activity", ""),
                max_chars=200,
            )
            updated_declaration = st.text_area(
                "Declaration text*",
                value=active_event.get("declaration", DECLARATION_PLACEHOLDER),
                help="Shown on the sign-in form under the Declaration section.",
                height=120,
            )
            updated_description = st.text_area(
                "Notes (optional)",
                value=active_event.get("description", ""),
                height=100,
            )
            col_save, col_cancel = st.columns(2)
            save_details = col_save.form_submit_button("Save details", type="primary")
            cancel_edit = col_cancel.form_submit_button("Cancel editing")
        if save_details:
            required_fields = {
                "Education name": updated_name.strip(),
                "Education date": updated_date.strip(),
                "Location": updated_location.strip(),
                "Project activity": updated_project_activity.strip(),
                "Declaration text": updated_declaration.strip(),
            }
            missing = [label for label, value in required_fields.items() if not value]
            if missing:
                st.error(f"Please fill in: {', '.join(missing)}.")
            else:
                try:
                    update_event_details(
                        event_id,
                        updated_name.strip(),
                        updated_date.strip(),
                        updated_location.strip(),
                        updated_project_activity.strip(),
                        updated_project_type,
                        updated_declaration.strip(),
                        updated_description.strip(),
                    )
                    st.success("Session details saved.")
                    st.session_state["active_event_id"] = event_id
                    st.session_state.pop("editing_event_id", None)
                    st.rerun()
                except Exception as exc:  # noqa: BLE001
                    st.error(f"Could not update session details: {exc}")
        elif cancel_edit:
            st.session_state.pop("editing_event_id", None)
            st.info("Editing cancelled.")
            st.rerun()
    else:
        st.info("Select a session and click edit to modify its details.")
        if st.button("Edit selected session", key=f"edit_{event_id}"):
            st.session_state["editing_event_id"] = event_id
            st.rerun()

    st.subheader("Replace attendee list")
    uploaded_file = st.file_uploader(
        "Upload CSV with columns: attendee_id, name, company, email", type=["csv"]
    )
    if uploaded_file is not None:
        try:
            new_df = pd.read_csv(uploaded_file, dtype=str).fillna("")
            replace_attendees(event_id, new_df)
            st.success("Attendee list updated successfully.")
            st.rerun()
        except Exception as exc:  # noqa: BLE001 (streamlit feedback)
            st.error(f"Could not import attendee list: {exc}")

    st.subheader("Downloads")
    signins_csv = signins.to_csv(index=False).encode("utf-8")
    st.download_button(
        "Download sign-ins as CSV",
        data=signins_csv,
        file_name=f"{event_id}-signins-{datetime.now().strftime('%Y%m%d-%H%M%S')}.csv",
        mime="text/csv",
    )

    template_path = ASSETS_DIR / "SignatureList.docx"
    if template_path.exists():
        if signins.empty:
            st.caption("Collect sign-ins to enable the signature sheet download.")
        else:
            try:
                doc_filename, doc_bytes = generate_signature_document(event_id)
                st.download_button(
                    "Download signature sheet (DOCX)",
                    data=doc_bytes,
                    file_name=doc_filename,
                    mime=(
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    ),
                )
            except Exception as exc:  # noqa: BLE001
                st.error(f"Unable to generate signature sheet: {exc}")
    else:
        st.caption(
            "SignatureList.docx not found in assets directory. Upload it to enable document export."
        )

    if st.button(
        "Clear sign-in records",
        help="Use at the start of a new education session.",
    ):
        pd.DataFrame(columns=SIGNIN_COLUMNS).to_csv(signin_file(event_id), index=False)
        load_signins.clear()
        st.success("Sign-in records cleared.")
        st.rerun()

    st.subheader("Create new education session")
    with st.form("create_event_form"):
        new_event_name = st.text_input(
            "Education name*", placeholder="e.g. Python Basics workshop"
        )
        new_event_project_type = st.selectbox(
            "Project template*",
            PROJECT_TYPES,
            help="Select the project this education belongs to.",
        )
        template_preview_new = get_project_template(new_event_project_type)
        st.caption(
            f"Preview: {template_preview_new.get('tagline', '')} "
            f"(image: {template_preview_new.get('image', 'N/A')})"
        )
        col_new_date, col_new_location = st.columns(2)
        new_event_date = col_new_date.text_input(
            "Education date*", placeholder="YYYY-MM-DD"
        )
        new_event_location = col_new_location.text_input(
            "Location*", placeholder="City / venue"
        )
        new_event_project_activity = st.text_input(
            "Project activity*", placeholder="Project or activity name"
        )
        new_event_description = st.text_area(
            "Notes (optional)",
            placeholder="Internal notes about this education (not shown to attendees).",
            height=80,
        )
        create_pressed = st.form_submit_button("Create session")
    if create_pressed:
        required_values = {
            "Education name": new_event_name.strip(),
            "Education date": new_event_date.strip(),
            "Location": new_event_location.strip(),
            "Project activity": new_event_project_activity.strip(),
            "Project template": new_event_project_type.strip(),
        }
        missing = [label for label, value in required_values.items() if not value]

        if missing:
            st.error(f"Please provide: {', '.join(missing)}.")
        else:
            try:
                record = create_event(
                    new_event_name.strip(),
                    new_event_date.strip(),
                    new_event_location.strip(),
                    new_event_project_activity.strip(),
                    new_event_project_type.strip(),
                    DECLARATION_PLACEHOLDER,
                    new_event_description.strip(),
                )
                st.success(
                    f"Created new session: {record['name'] or record['event_id']}"
                )
                st.session_state["active_event_id"] = record["event_id"]
                st.session_state["reset_attendee_search"] = True
                st.session_state.pop("selected_attendee", None)
                st.session_state.pop("signature_canvas_key", None)
                st.rerun()
            except Exception as exc:  # noqa: BLE001
                st.error(f"Could not create session: {exc}")

    if st.button("Log out"):
        st.session_state["admin_authenticated"] = False
        st.success("You have been logged out.")
        st.rerun()


def main() -> None:
    ensure_storage()
    st.set_page_config(
        page_title="Education Sign-In", page_icon=":memo:", layout="wide"
    )

    if "admin_authenticated" not in st.session_state:
        st.session_state["admin_authenticated"] = False

    events_df = load_events()
    event_records = events_df.to_dict("records")
    event_lookup = {row["event_id"]: row for row in event_records}

    if not event_records:
        st.error("No education sessions available. Please create one in the admin panel.")
        return

    default_event_id = next(
        (
            row["event_id"]
            for row in event_records
            if str(row.get("is_default", "")).lower() == "true"
        ),
        event_records[0]["event_id"],
    )
    active_event_id = st.session_state.get("active_event_id", default_event_id)
    if active_event_id not in event_lookup:
        active_event_id = default_event_id
    st.session_state["active_event_id"] = active_event_id

    event_ids = [row["event_id"] for row in event_records]
    default_index = event_ids.index(active_event_id)

    def sidebar_label(eid: str) -> str:
        label = event_lookup[eid].get("name") or eid
        if str(event_lookup[eid].get("is_default", "")).lower() == "true":
            label += " (default)"
        return label

    selected_event_id = st.sidebar.selectbox(
        "Education session",
        event_ids,
        index=default_index,
        format_func=sidebar_label,
        key="event_selector",
    )
    if selected_event_id != active_event_id:
        st.session_state["active_event_id"] = selected_event_id
        st.session_state.pop("selected_attendee", None)
        st.session_state["reset_attendee_search"] = True
        st.session_state.pop("signature_canvas_key", None)
    active_event = event_lookup[selected_event_id]

    if active_event.get("description"):
        st.sidebar.caption(active_event["description"])

    page = st.sidebar.radio("Navigation", ["Sign In", "Admin"])
    if page == "Admin":
        if st.session_state.get("admin_authenticated", False):
            admin_page(selected_event_id, events_df, active_event)
        else:
            admin_login_page()
    else:
        sign_in_page(selected_event_id, active_event)


if __name__ == "__main__":
    main()
